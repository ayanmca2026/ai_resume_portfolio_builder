import os
from docx import Document
from docx.shared import Pt, Inches

def generate_resume_docx(resume_data: dict, output_path: str) -> str:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    document = Document()
    
    # Setting minimal margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    name = resume_data.get("name", "User Name")
    document.add_heading(name, 0)
    
    if resume_data.get("summary"):
        document.add_heading("Professional Summary", level=1)
        document.add_paragraph(resume_data["summary"])
        
    if resume_data.get("skills"):
        document.add_heading("Skills", level=1)
        document.add_paragraph(", ".join(resume_data["skills"]))
        
    if resume_data.get("experience"):
        document.add_heading("Experience", level=1)
        for exp in resume_data["experience"]:
            p = document.add_paragraph()
            p.add_run(exp.get("role", "")).bold = True
            p.add_run(f" at {exp.get('company', '')} ")
            p.add_run(f"({exp.get('duration', '')})").italic = True
            if exp.get('responsibilities'):
                document.add_paragraph(exp.get('responsibilities'), style='List Bullet')
                
    if resume_data.get("education"):
        document.add_heading("Education", level=1)
        for edu in resume_data["education"]:
            p = document.add_paragraph()
            p.add_run(edu.get("degree", "")).bold = True
            p.add_run(f", {edu.get('university', '')} ({edu.get('graduation_year', '')})")

    document.save(output_path)
    return output_path
